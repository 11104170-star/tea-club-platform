from __future__ import annotations

from collections import Counter
from io import BytesIO
from pathlib import Path
from typing import BinaryIO

import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE_PATH = PROJECT_ROOT / "assets" / "成果書模板_已標記.docx"

SCORE_VALUES = {"1", "2", "3", "4", "5"}


def read_questionnaire(uploaded_file) -> pd.DataFrame:
    ext = Path(uploaded_file.name).suffix.lower()
    uploaded_file.seek(0)

    if ext == ".xlsx":
        return pd.read_excel(uploaded_file)

    for encoding in ("utf-8-sig", "utf-8", "cp950", "big5"):
        try:
            uploaded_file.seek(0)
            return pd.read_csv(uploaded_file, encoding=encoding)
        except UnicodeDecodeError:
            continue

    raise ValueError("CSV 編碼無法辨識，請改存為 UTF-8 或 Excel 檔後再上傳。")


def analyze_questionnaire(df: pd.DataFrame) -> str:
    df = df.dropna(axis=1, how="all")
    total = len(df)

    if total == 0:
        return "本次問卷尚無有效回覆。"

    score_questions = []
    text_questions = []

    for column in df.columns:
        values = df[column].dropna().astype(str).str.strip()
        non_empty = values[values != ""]

        if len(non_empty) > 0 and non_empty.isin(SCORE_VALUES).all():
            score_questions.append(column)
        else:
            text_questions.append(column)

    lines = [f"本次問卷有效回覆數：{total} 份", ""]

    if score_questions:
        lines.append("量表題統計（1-5 分）：")
        for index, question in enumerate(score_questions, start=1):
            values = df[question].dropna().astype(str).str.strip()
            counter = Counter(values[values != ""])
            parts = []

            for score in sorted(SCORE_VALUES, reverse=True):
                count = counter.get(score, 0)
                percent = count / total * 100
                parts.append(f"{score} 分：{count} 份（{percent:.1f}%）")

            lines.append(f"{index}. {question}")
            lines.append("、".join(parts))

    if text_questions:
        if score_questions:
            lines.append("")
        lines.append("文字題回覆整理：")
        start_index = len(score_questions) + 1

        for index, question in enumerate(text_questions, start=start_index):
            counter = Counter()
            blank_count = 0

            for value in df[question]:
                if pd.isna(value) or str(value).strip() == "":
                    blank_count += 1
                else:
                    counter[str(value).strip()] += 1

            lines.append(f"{index}. {question}")
            for text, count in counter.items():
                lines.append(f"- {text}（{count} 份）")
            lines.append(f"- 未填答（{blank_count} 份）")

    return "\n".join(lines)


def set_font(paragraph) -> None:
    for run in paragraph.runs:
        run.font.name = "DFKai-SB"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
        run.font.size = Pt(11)


def iter_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def replace_text(doc, replacements: dict[str, str]) -> None:
    for paragraph in iter_paragraphs(doc):
        original_text = paragraph.text
        new_text = original_text

        for key, value in replacements.items():
            new_text = new_text.replace(key, value or "")

        if new_text != original_text:
            paragraph.text = new_text
            set_font(paragraph)


def image_stream(uploaded_file) -> BinaryIO | None:
    if uploaded_file is None:
        return None

    return BytesIO(uploaded_file.getvalue())


def insert_images(doc, images: dict[str, object]) -> None:
    for paragraph in iter_paragraphs(doc):
        for key, uploaded_file in images.items():
            if key in paragraph.text:
                paragraph.text = ""
                stream = image_stream(uploaded_file)

                if stream is not None:
                    run = paragraph.add_run()
                    run.add_picture(stream, width=Inches(2.5))


def build_report(
    *,
    template_file,
    questionnaire_file,
    fields: dict[str, str],
    images: dict[str, object],
) -> tuple[BytesIO, str]:
    if questionnaire_file is None:
        result_text = "未上傳問卷資料。"
    else:
        df = read_questionnaire(questionnaire_file)
        result_text = analyze_questionnaire(df)

    template_source = template_file if template_file is not None else str(DEFAULT_TEMPLATE_PATH)
    doc = Document(template_source)

    replacements = {
        "{{填寫日期}}": fields.get("fill_date", ""),
        "{{活動名稱}}": fields.get("activity_name", ""),
        "{{活動地點}}": fields.get("activity_place", ""),
        "{{活動日期}}": fields.get("activity_date", ""),
        "{{活動負責人}}": fields.get("activity_leader", ""),
        "{{連絡電話}}": fields.get("phone", ""),
        "{{參加人數}}": fields.get("activity_people", ""),
        "{{問卷分析結果}}": result_text,
        "{{活動檢討}}": fields.get("activity_review", ""),
        "{{照片1說明}}": fields.get("photo1_desc", ""),
        "{{照片2說明}}": fields.get("photo2_desc", ""),
        "{{照片3說明}}": fields.get("photo3_desc", ""),
        "{{指導老師評語}}": fields.get("teacher_comment", ""),
    }

    replace_text(doc, replacements)

    insert_images(
        doc,
        {
            "{{活動流程照片}}": images.get("flow_photo"),
            "{{大合照 }}": images.get("group_photo"),
            "{{大合照}}": images.get("group_photo"),
            "{{照片1 }}": images.get("photo1"),
            "{{照片1}}": images.get("photo1"),
            "{{照片2 }}": images.get("photo2"),
            "{{照片2}}": images.get("photo2"),
            "{{照片3 }}": images.get("photo3"),
            "{{照片3}}": images.get("photo3"),
        },
    )

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return output, result_text
